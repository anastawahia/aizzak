"""DOCX parsing (ported from alpha `rag/parsers/text_parser.py` —
`RAGFlowTextParser._extract_docx_raw`; docs/rag-indexing-plan.md step 4, item
`P-08`).

What is ported is the **algorithm and its calibration**: the heading ladder,
the caption/section two-level hierarchy, the paragraph-merge rule, and the
2000-character block ceiling. Alpha's structure — environment-variable
configuration (`DOCX_SAVE_CHUNKS_TO_DISK`), disk dumping, the stateful
`self.logs` accumulator and its `print_logs` — stays behind.

**Headings: Word styles first, the text heuristic as the fallback** (decision
س-09 = أ). A paragraph styled `Heading 1..9` / `Title` / `Subtitle` is a
heading because the *document* says so — a structural signal, and therefore a
language-neutral one. Alpha has this check (`is_heading_style`) but consults it
alongside a heuristic built on English keywords and Title-Case rules, which do
not fire on an Arabic document at all (fact ح-١٤). Here the style wins outright
and the heuristic only answers for paragraphs that carry no heading style.
Detecting the heading is what makes "section such-and-such" citable in a file
that has no page numbers, which is why it is worth this much care.

**Document order comes from `iter_inner_content()`.** Paragraphs and tables
interleave, and a table's meaning depends on the heading above it, so both have
to be walked in one sequence. Alpha reconstructs that sequence by comparing
`_element.getparent().index(...)` between `doc.paragraphs` and `doc.tables`;
python-docx ≥ 1.1 exposes the same sequence directly. Same decision, supported
API instead of XML internals.

`ParsedChunk.order` is `position_in_doc * 1000 + part_index` — a deterministic
STRUCTURAL ordinal (the port's contract): it is derived from where the block
sits in the document, never from how many blocks happened to survive the merge,
so a table and the paragraph after it keep their relative order and a block
split for length keeps its parts adjacent.

Two divergences from alpha, both deliberate:

- **A block's `position_in_doc` is its own first paragraph's**, not that of the
  heading that ended it. Alpha stamps the emitted block with the position of
  the *next* heading (and drops the field entirely on the final flush), which
  makes the field non-monotonic — unusable as the ordering signal step 10
  (`P-17`) consumes it as.
- **The heading text is prepended to the block text**, as alpha does, and also
  carried in ``metadata["title"]``. `pdf_tables.py` deliberately does the
  opposite with its captions; the difference is real: a PDF caption is a text
  block the text pass indexes anyway, while a DOCX heading is consumed by this
  merge and would otherwise never reach the index at all.

Not in this module: OCR of the images embedded in a DOCX — landed in step 5
(`P-09`) but in `image_ocr.py`, which reads `word/media/` straight out of the
OOXML archive, because `python-docx` models a picture as an inline shape and
never hands over its raster; and table-row explosion into one node per row
(step 7 / `P-13`) — a table stays one coarse chunk here, as it does for Excel
and PDF.
"""

from __future__ import annotations

import io
import json
import re
from typing import NamedTuple

from docx import Document
from docx.document import Document as WordDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.framework.observability import get_logger
from app.framework.types import Json
from app.modules.knowledge.adapters.parsers.text_plain import split_long_text
from app.modules.knowledge.ports.content_extractor import ParsedChunk, ParsedChunkKind

log = get_logger(__name__)

# alpha `DOCX_MAX_CHUNK_CHARS` (default 2000) — the same ceiling `text_plain.py`
# already reuses for plain text.
MAX_CHUNK_CHARS = 2000

# A heading is repeated at the top of every part of a split block, so the room
# left for content shrinks by its length; never below this floor, or a long
# heading would shred the body into fragments (alpha's `max(..., 200)`).
_MIN_SPLIT_CHARS = 200

# order = position_in_doc * STRIDE + part_index. A block would need 1000 parts
# (~2 MB of text under one heading) to reach the next item's range; the part
# index is clamped instead, which at worst ties the tail of such a block — a
# stable sort keeps those in order — and can never reorder two items.
_ITEM_ORDER_STRIDE = 1_000

# alpha's `_classify_text` ratio bands, with `docx` as the base label.
_STRUCTURED_RATIO = 0.6
_SEMI_STRUCTURED_RATIO = 0.2

# A table needs a header row and at least one data row to be a table (alpha).
_MIN_TABLE_ROWS = 2

# Word style names that mark a heading outright (checked lower-cased; the
# `Heading 1..9` family is matched by prefix).
_HEADING_STYLE_PREFIX = "heading"
_HEADING_STYLE_NAMES = frozenset({"title", "subtitle"})

# Sentence-ending punctuation, Arabic included: a real heading almost never
# ends with it, which is what keeps ordinary sentences out of the ladder.
_SENTENCE_ENDINGS = ".!?،؛"

# A label heading ("Purpose:") is short; past this it is a sentence that
# happens to end in a colon.
_LABEL_MAX_CHARS = 120
# Classic short headings ("Job Analysis") — capitalized and this short.
_SHORT_HEADING_MAX_WORDS = 4
# Longer Title-Case / ALL-CAPS headings ("Internal and External Communication
# Guidelines") are still accepted up to this length.
_TITLE_CASE_MAX_WORDS = 9
_TITLE_CASE_MIN_WORDS = 2
# A pipe-delimited line with this many filled cells is a data row, not prose.
_TABULAR_MIN_CELLS = 2

# Small words that stay lowercase in Title Case — they must NOT count against a
# line being a heading (alpha `_TITLE_STOPWORDS`).
_TITLE_STOPWORDS = frozenset(
    {
        "a", "an", "the", "and", "or", "of", "for", "to", "in", "on", "at",
        "by", "with", "from", "as", "but", "nor", "per", "via", "vs", "&",
    }
)  # fmt: skip

# Known heading words, matched after stripping leading numbering and a trailing
# colon (alpha `_HEADING_KEYWORDS`). English-only, as in alpha: decision س-09
# took option (أ), which relies on Word styles for language neutrality and
# keeps this list as the fallback for unstyled English documents. Adding the
# Arabic equivalents was option (ب) and was not taken.
_HEADING_KEYWORDS = frozenset(
    {
        "purpose", "scope", "objective", "objectives", "responsibilities",
        "responsibility", "introduction", "definitions", "definition",
        "procedure", "procedures", "claim procedure", "policy", "process",
        "overview", "references", "guidelines", "general", "eligibility",
        "roles and responsibilities", "annex", "appendix", "notes",
    }
)  # fmt: skip

_NUMBERED_HEADING = re.compile(r"^\d+(\.\d+)+([.\):])?\s+\S")
_SECTIONED_HEADING = re.compile(
    r"^(section|article|chapter|annex|appendix|part)\s+[\divxlc]+\b", re.IGNORECASE
)
_LEADING_NUMBERING = re.compile(r"^\d+([.\)]|\.\d+)*\s*")
_BULLET_NUMBERING = re.compile(r"^\d+[\.\)]")
_LABEL_WITH_BODY = re.compile(r":\s+\S")
# The en/em dashes are bullet glyphs Word actually emits, not punctuation:
_BULLET_MARKERS = ("-", "•", "*", "▪", "–", "—")  # noqa: RUF001


class _Item(NamedTuple):
    """One paragraph or table, in document order.

    ``paragraph_number`` is the paragraph's ordinal among the document's
    paragraphs (empty ones included, as in alpha) and ``table_index`` the
    table's among its tables; each is `None` for the other kind.
    """

    paragraph: Paragraph | None
    table: Table | None
    position_in_doc: int
    paragraph_number: int | None
    table_index: int | None

    @property
    def text(self) -> str:
        return (self.paragraph.text or "").strip() if self.paragraph is not None else ""


class _Block(NamedTuple):
    """A heading with the paragraphs merged under it (either part may stand
    alone: a heading with no body, or body text before any heading)."""

    title: str | None
    content: str
    paragraph_number: int | None
    position_in_doc: int


def parse_docx(data: bytes) -> list[ParsedChunk]:
    """Extract merged text blocks and structured tables from a ``.docx`` file.

    Raises whatever `python-docx` raises on bytes that are not a readable Word
    document — `adapters/parsers/extractor.py` translates that into a
    `ValidationError`, the same policy `excel.py` follows. There is no
    per-paragraph fault tolerance to mirror (alpha has none either): unlike a
    PDF, a DOCX either opens or does not.
    """
    document = Document(io.BytesIO(data))
    items = _document_sequence(document)
    return [*_text_chunks(items), *_table_chunks(items)]


def classify_docx(*, table_count: int, text_count: int, image_count: int = 0) -> str:
    """File-type classification by table/text ratio (alpha `_classify_text`,
    with ``docx`` as its ``base_label``).

    ``image_count`` (plan step 5 / `P-09`) decides only the empty case: a
    document whose words all live inside pictures is not empty, and the ratio
    bands below are about the text/table mix, which an image does not join.
    """
    if table_count == 0 and text_count == 0:
        return "docx_images" if image_count else "docx_empty"
    if table_count > 0 and text_count == 0:
        return "docx_structured_text"
    if table_count == 0 and text_count > 0:
        return "docx_unstructured_text"

    ratio = table_count / max(table_count + text_count, 1)
    if ratio > _STRUCTURED_RATIO:
        return "docx_structured_text"
    if ratio > _SEMI_STRUCTURED_RATIO:
        return "docx_semi_structured_text"
    return "docx_unstructured_text"


def _document_sequence(document: WordDocument) -> list[_Item]:
    """Paragraphs and tables in one document-ordered sequence."""
    items: list[_Item] = []
    paragraph_number = 0
    table_index = 0
    for position, content in enumerate(document.iter_inner_content()):
        if isinstance(content, Table):
            items.append(_Item(None, content, position, None, table_index))
            table_index += 1
        else:
            items.append(_Item(content, None, position, paragraph_number, None))
            paragraph_number += 1
    return items


# --------------------------------------------------------------------------- #
# Text: heading detection, merging, splitting                                 #
# --------------------------------------------------------------------------- #
def _text_chunks(items: list[_Item]) -> list[ParsedChunk]:
    return [chunk for block in _merge_blocks(items) for chunk in _block_to_chunks(block)]


def _merge_blocks(items: list[_Item]) -> list[_Block]:
    """Fold paragraphs into one block per heading (alpha's merge loop).

    A heading closes the block before it and opens the next; bullets and body
    text accumulate under the open heading. A paragraph with no heading above
    it stands alone rather than joining an unrelated neighbour.
    """
    blocks: list[_Block] = []
    title: str | None = None
    buffer: list[str] = []
    first_paragraph: int | None = None
    first_position: int | None = None

    def flush() -> None:
        if title is None and not buffer:
            return
        blocks.append(
            _Block(title, "\n".join(buffer).strip(), first_paragraph, first_position or 0)
        )

    for index, item in enumerate(items):
        text = item.text
        if item.paragraph is None or not text:
            continue

        if _opens_a_section(item, items, index):
            flush()
            title, buffer = text, []
            first_paragraph, first_position = item.paragraph_number, item.position_in_doc
            continue

        if first_paragraph is None:
            first_paragraph = item.paragraph_number
        if first_position is None:
            first_position = item.position_in_doc

        if _is_bullet(text) or title is not None:
            buffer.append(text)
            continue

        blocks.append(_Block(None, text, item.paragraph_number, item.position_in_doc))
        first_paragraph = first_position = None

    flush()
    return blocks


def _opens_a_section(item: _Item, items: list[_Item], index: int) -> bool:
    """Whether this paragraph starts a new section.

    A weak heading (one ending in a colon) directly above tabular text is a
    **caption** for that data, not a section of its own: treating it as body
    keeps the rows under the descriptive heading they belong to instead of
    starting a detached, context-free block. The rule is typographic, not a
    list of phrases, so it holds across documents and languages.
    """
    text = item.text
    if _is_weak_heading(text) and _is_title(text) and _looks_tabular(_next_text(items, index)):
        return False
    return _is_heading_style(item.paragraph) or _is_title(text)


def _next_text(items: list[_Item], index: int) -> str:
    """Text of the next non-empty paragraph after ``index``."""
    for item in items[index + 1 :]:
        if item.paragraph is not None and item.text:
            return item.text
    return ""


def _is_heading_style(paragraph: Paragraph | None) -> bool:
    """Whether Word itself marks this paragraph as a heading.

    The most reliable signal in a Word document, and the only one that says
    nothing about the language the document is written in (decision س-09).
    """
    if paragraph is None:
        return False
    try:
        style = paragraph.style
        name = (style.name or "").strip().lower() if style is not None else ""
    except Exception:  # a style id the document does not define
        return False
    return name.startswith(_HEADING_STYLE_PREFIX) or name in _HEADING_STYLE_NAMES


def _is_title(text: str) -> bool:
    """The text heuristic (alpha `is_title`) — the fallback for a paragraph
    carrying no heading style.

    Alpha's single ladder of returns is split into the named predicates below;
    the rules, their order and their thresholds are unchanged.
    """
    normalized = re.sub(r"\s+", " ", text or "").strip()
    if not normalized or normalized[-1] in _SENTENCE_ENDINGS or "|" in normalized:
        return False
    return (
        _is_label_heading(normalized)
        or _is_numbered_heading(normalized)
        or _is_keyword_heading(normalized)
        or _is_capitalized_heading(normalized)
    )


def _is_label_heading(text: str) -> bool:
    """A short label ending in a colon ("Purpose:", "Responsibilities:").

    Rejects "HR & Admin: the team is responsible for:" — a colon with text
    after it is a paragraph lead-in, not a heading.
    """
    return (
        text.endswith(":") and len(text) <= _LABEL_MAX_CHARS and not _LABEL_WITH_BODY.search(text)
    )


def _is_numbered_heading(text: str) -> bool:
    """ "5.1 Internal ...", "2.3.1 Scope", "Annex 1 ...", "Section IV ..."."""
    return bool(_NUMBERED_HEADING.match(text) or _SECTIONED_HEADING.match(text))


def _is_keyword_heading(text: str) -> bool:
    """A known heading word, once leading numbering and a trailing colon are
    stripped ("3.1 Scope:" -> "scope")."""
    core = _LEADING_NUMBERING.sub("", text).rstrip(":").strip().lower()
    return core in _HEADING_KEYWORDS


def _is_capitalized_heading(text: str) -> bool:
    """Short capitalized lines, ALL-CAPS lines, and Title-Case lines."""
    words = text.split()
    count = len(words)
    if 1 <= count <= _SHORT_HEADING_MAX_WORDS and text[0].isupper():
        return True
    if not (_TITLE_CASE_MIN_WORDS <= count <= _TITLE_CASE_MAX_WORDS):
        return False

    letters = [c for c in text if c.isalpha()]
    if letters and all(c.isupper() for c in letters):
        return True
    alpha_words = [w for w in words if any(c.isalpha() for c in w)]
    significant = [w for w in alpha_words if w.lower() not in _TITLE_STOPWORDS]
    return bool(significant) and text[0].isupper() and all(w[0].isupper() for w in significant)


def _is_bullet(text: str) -> bool:
    stripped = text.lstrip()
    return stripped.startswith(_BULLET_MARKERS) or bool(_BULLET_NUMBERING.match(stripped))


def _is_weak_heading(text: str) -> bool:
    """A "weak" heading is a caption/lead-in label: it ends with a colon. A
    "strong" one is a descriptive section name. The distinction builds the
    two-level hierarchy a table's title uses."""
    return re.sub(r"\s+", " ", text or "").strip().endswith(":")


def _looks_tabular(text: str) -> bool:
    """A pipe-delimited data row, e.g. ``Class A1 | 150000 | A``."""
    return "|" in text and len([c for c in text.split("|") if c.strip()]) >= _TABULAR_MIN_CELLS


def _block_to_chunks(block: _Block) -> list[ParsedChunk]:
    """One block into one chunk, or several when it exceeds `MAX_CHUNK_CHARS`.

    The heading is repeated at the top of every part: each part is retrieved on
    its own, and a part that opens mid-argument with no section name above it
    is a part no reader can place.
    """
    content = block.content.strip()
    if not content:
        return []

    header = f"{block.title}\n" if block.title else ""
    full_text = f"{header}{content}"
    if len(full_text) <= MAX_CHUNK_CHARS:
        parts = [full_text]
    else:
        limit = max(MAX_CHUNK_CHARS - len(header), _MIN_SPLIT_CHARS)
        parts = [f"{header}{part}" for part in split_long_text(content, limit)]

    return [
        ParsedChunk(
            text=part,
            order=_order_of(block.position_in_doc, part_index),
            kind=ParsedChunkKind.TEXT,
            metadata=_text_metadata(block, part_index=part_index, part_count=len(parts)),
        )
        for part_index, part in enumerate(parts)
    ]


def _order_of(position_in_doc: int, part_index: int) -> int:
    return position_in_doc * _ITEM_ORDER_STRIDE + min(part_index, _ITEM_ORDER_STRIDE - 1)


def _text_metadata(block: _Block, *, part_index: int, part_count: int) -> Json:
    metadata: Json = {
        "paragraph_number": block.paragraph_number,
        "position_in_doc": block.position_in_doc,
        "part_index": part_index,
        "part_count": part_count,
        "section_type": "section" if block.title else "paragraph",
        "chunk_type": "docx_text",
        "source_ext": ".docx",
    }
    if block.title:
        metadata["title"] = block.title
    return metadata


# --------------------------------------------------------------------------- #
# Tables                                                                      #
# --------------------------------------------------------------------------- #
def _table_chunks(items: list[_Item]) -> list[ParsedChunk]:
    """Every table with the nearest heading above it as its title.

    The title is a two-level breadcrumb — the nearest *descriptive* section
    heading plus the immediate *caption* — so a generic caption ("Quick
    Reference Table:") never shadows the policy the table belongs to. Without a
    descriptive anchor a table's rows retrieve as a wall of numbers.
    """
    chunks: list[ParsedChunk] = []
    section: str | None = None
    caption: str | None = None

    for item in items:
        if item.table is None:
            text = item.text
            if text and (_is_heading_style(item.paragraph) or _is_title(text)):
                if _is_weak_heading(text):
                    caption = text
                else:
                    section, caption = text, None
            continue

        chunk = _table_to_chunk(item, _breadcrumb(section, caption))
        if chunk is not None:
            chunks.append(chunk)
    return chunks


def _breadcrumb(section: str | None, caption: str | None) -> str:
    """``Attendance Policy — Quick Reference Table:``; either part may be
    missing, and a caption identical to its section is not repeated."""
    parts = [part for part in (section, caption) if part]
    return " — ".join(dict.fromkeys(parts))


def _table_to_chunk(item: _Item, title: str) -> ParsedChunk | None:
    """One table into one coarse chunk, or `None` when it has no data rows."""
    if item.table is None:
        return None
    rows = _table_rows(item.table)
    if len(rows) < _MIN_TABLE_ROWS:
        log.info("docx.table_skipped", extra={"table_index": item.table_index, "rows": len(rows)})
        return None

    headers = [cell.strip() or f"Column_{index + 1}" for index, cell in enumerate(rows[0])]
    records = [dict(zip(headers, _fit(row, len(headers)), strict=True)) for row in rows[1:]]

    metadata: Json = {
        "table_index": item.table_index,
        "position_in_doc": item.position_in_doc,
        "headers": headers,
        "num_rows": len(records),
        "num_cols": len(headers),
        "section_type": "structured_table",
        "chunk_type": "docx_table",
        "source_ext": ".docx",
    }
    if title:
        metadata["title"] = title

    return ParsedChunk(
        text=json.dumps({"headers": headers, "rows": records}, ensure_ascii=False),
        order=_order_of(item.position_in_doc, 0),
        kind=ParsedChunkKind.TABLE,
        metadata=metadata,
    )


def _table_rows(table: Table) -> list[list[str]]:
    """The table's non-empty rows as stripped cell text."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [(cell.text or "").strip() for cell in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def _fit(row: list[str], width: int) -> list[str]:
    """Truncate or pad a data row to the header's width, so every record has
    the same keys (alpha normalizes the same way)."""
    return row[:width] + [""] * max(0, width - len(row))
