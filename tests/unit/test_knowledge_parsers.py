"""Unit tests for the knowledge document parsers (3.k1 — docs/migration/refs/
parsers.md). No committed binary fixtures: PDF/Excel/image inputs are
synthesized in-test via PyMuPDF/openpyxl/Pillow. Plain unit tests — no
``integration`` marker, no Docker — but they require the ``parsers`` optional
dependency group (``pip install -e ".[dev,parsers]"``)."""

from __future__ import annotations

import io
import json
import shutil
import zipfile
from collections.abc import Callable

import fitz
import openpyxl
import pandas as pd
import pytesseract
import pytest
from docx import Document as WordDocument
from docx.document import Document as WordDocumentObject
from openpyxl.drawing.image import Image as WorksheetImage
from PIL import Image

from app.framework.context.execution_context import ExecutionContext
from app.framework.errors import UnsupportedTypeError, ValidationError
from app.framework.ports.embedding_provider import EmbeddingResult
from app.framework.ports.vector_store import VectorPoint
from app.framework.settings.settings import Limits
from app.framework.types import Json
from app.modules.knowledge.adapters.parsers import docx as docx_parser
from app.modules.knowledge.adapters.parsers import (
    excel,
    image_ocr,
    json_doc,
    pdf_tables,
    pdf_text,
    text_plain,
)
from app.modules.knowledge.adapters.parsers import extractor as extractor_module
from app.modules.knowledge.adapters.parsers.extractor import DocumentContentExtractor
from app.modules.knowledge.application.indexing import IndexDocument
from app.modules.knowledge.ports.content_extractor import (
    ContentExtractor,
    ParsedChunkKind,
    ParsedDocument,
)


# --------------------------------------------------------------------------- #
# json_doc — table patterns, classification, encoding                        #
# --------------------------------------------------------------------------- #
def test_parse_json_empty_object_yields_no_chunks() -> None:
    assert json_doc.parse_json(b"{}") == []


def test_parse_json_list_of_dict_table_pattern() -> None:
    data = [{"a": 1, "b": 2}, {"a": 3, "b": 4}]
    chunks = json_doc.parse_json(json.dumps(data).encode())

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.kind is ParsedChunkKind.TABLE
    assert chunk.order == 0
    assert chunk.metadata["source_type"] == "dict_rows"
    assert chunk.metadata["headers"] == ["a", "b"]
    assert json.loads(chunk.text)["rows"] == data


def test_parse_json_grid_list_of_str_table_pattern() -> None:
    data = ["Name\tAge", "Ali\t30", "Sara\t25"]
    chunks = json_doc.parse_json(json.dumps(data).encode())

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.kind is ParsedChunkKind.TABLE
    assert chunk.metadata["source_type"] == "structured_text"
    assert chunk.metadata["headers"] == ["Column_1", "Column_2"]
    assert chunk.metadata["num_rows"] == 3


def test_parse_json_kv_scalar_table_pattern() -> None:
    data = {"name": "Ali", "age": 30, "city": "Riyadh", "active": True}
    chunks = json_doc.parse_json(json.dumps(data).encode())

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.kind is ParsedChunkKind.TABLE
    assert chunk.metadata["headers"] == ["Key", "Value"]
    assert chunk.metadata["num_rows"] == 4


def test_parse_json_embedded_text_table_inside_string_value() -> None:
    # a table pattern nested INSIDE a dict value (exercises the recursive walk).
    data = {"report": "Name\tAge\nAli\t30\nSara\t25\nOmar\t40"}
    chunks = json_doc.parse_json(json.dumps(data).encode())

    assert len(chunks) == 1
    assert chunks[0].kind is ParsedChunkKind.TABLE
    assert chunks[0].metadata["path"] == "$.report"
    assert chunks[0].metadata["num_rows"] == 4


def test_parse_json_unstructured_text_only() -> None:
    data = {"title": "Hello", "body": "This is a plain description with some prose."}
    chunks = json_doc.parse_json(json.dumps(data).encode())

    assert len(chunks) == 2
    assert all(c.kind is ParsedChunkKind.JSON for c in chunks)
    assert [c.order for c in chunks] == [0, 1]

    n_tables = sum(1 for c in chunks if c.kind is ParsedChunkKind.TABLE)
    n_texts = sum(1 for c in chunks if c.kind is ParsedChunkKind.JSON)
    assert json_doc.classify_json(table_count=n_tables, text_count=n_texts) == "unstructured_json"


@pytest.mark.parametrize(
    ("table_count", "text_count", "expected"),
    [
        (0, 0, "empty_json"),
        (2, 0, "structured_json"),
        (0, 2, "unstructured_json"),
        (8, 2, "structured_json"),  # ratio 0.8 > 0.6
        (3, 7, "semi_structured_json"),  # ratio 0.3, > 0.2
        (1, 9, "unstructured_json"),  # ratio 0.1 <= 0.2
    ],
)
def test_classify_json_thresholds(table_count: int, text_count: int, expected: str) -> None:
    assert json_doc.classify_json(table_count=table_count, text_count=text_count) == expected


def test_decode_bytes_round_trips_utf16() -> None:
    text = "بيانات تجريبية"
    assert json_doc._decode_bytes(text.encode("utf-16")) == text


def test_decode_bytes_round_trips_cp1256() -> None:
    text = "مرحبا"
    assert json_doc._decode_bytes(text.encode("cp1256")) == text


def test_parse_json_never_raises_on_undecodable_looking_bytes() -> None:
    # the candidate chain always has a lossless last resort (latin-1 / utf-8
    # with replacement) — malformed-looking bytes must never raise.
    chunks = json_doc.parse_json(b'{"note": "caf\xe9"}')
    assert isinstance(chunks, list)


# --------------------------------------------------------------------------- #
# text_plain — encoding, clean_text, coarse split, Arabic punctuation         #
# --------------------------------------------------------------------------- #
def test_parse_text_basic_ascii() -> None:
    chunks = text_plain.parse_text(b"Hello world, this is a test.", ".txt")

    assert len(chunks) == 1
    assert chunks[0].order == 0
    assert chunks[0].kind is ParsedChunkKind.TEXT
    assert chunks[0].metadata["source_ext"] == ".txt"
    assert "Hello world" in chunks[0].text


def test_parse_text_empty_bytes_yields_no_chunks() -> None:
    assert text_plain.parse_text(b"", ".txt") == []


@pytest.mark.parametrize("ext", [".txt", ".md", ".csv"])
def test_parse_text_uniform_across_routed_extensions(ext: str) -> None:
    chunks = text_plain.parse_text(b"col1,col2\nval1,val2", ext)
    assert len(chunks) == 1
    assert chunks[0].kind is ParsedChunkKind.TEXT
    assert chunks[0].metadata["source_ext"] == ext


def test_clean_text_strips_page_numbers_and_noise() -> None:
    cleaned = text_plain.clean_text("Page 12\n----- noise ----- \n\n\nReal content")
    assert "Page 12" not in cleaned
    assert "Real content" in cleaned


def test_split_long_text_keeps_short_text_as_single_piece() -> None:
    assert text_plain.split_long_text("short text", max_chars=2000) == ["short text"]


def test_split_long_text_splits_oversized_line_on_arabic_punctuation() -> None:
    line = "جملة أولى، جملة ثانية؛ جملة ثالثة"
    pieces = text_plain.split_long_text(line, max_chars=15)
    assert pieces == ["جملة أولى،", "جملة ثانية؛", "جملة ثالثة"]


def test_parse_text_deterministic_split_index_as_order() -> None:
    long_text = "\n\n".join(f"Paragraph number {i} with some filler content." for i in range(200))
    chunks = text_plain.parse_text(long_text.encode(), ".md")

    assert len(chunks) > 1
    assert [c.order for c in chunks] == list(range(len(chunks)))
    assert all(c.metadata["split_index"] == c.order for c in chunks)


# --------------------------------------------------------------------------- #
# excel — sheet segmentation, _split_large_table boundary, empty workbook     #
# --------------------------------------------------------------------------- #
def _workbook_bytes(build: Callable[[openpyxl.Workbook], None]) -> bytes:
    wb = openpyxl.Workbook()
    build(wb)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_parse_excel_emits_one_chunk_per_sheet_table() -> None:
    def build(wb: openpyxl.Workbook) -> None:
        ws = wb.active
        ws.title = "Data"
        ws.append(["ID", "Value"])
        for i in range(3):
            ws.append([i, i * 2])

    chunks = excel.parse_excel(_workbook_bytes(build))

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.kind is ParsedChunkKind.TABLE
    assert chunk.order == 0
    assert chunk.metadata["sheet_name"] == "Data"
    assert chunk.metadata["num_rows"] == 3
    assert json.loads(chunk.text)["headers"] == ["ID", "Value"]


def test_parse_excel_splits_large_table_at_max_rows_per_chunk_boundary() -> None:
    def build(wb: openpyxl.Workbook) -> None:
        ws = wb.active
        ws.title = "Big"
        ws.append(["ID", "Value"])
        for i in range(excel._MAX_ROWS_PER_CHUNK + 1):
            ws.append([i, i * 2])

    chunks = excel.parse_excel(_workbook_bytes(build))

    assert len(chunks) == 2
    assert chunks[0].metadata["num_rows"] == excel._MAX_ROWS_PER_CHUNK
    assert chunks[1].metadata["num_rows"] == 1
    assert chunks[0].metadata["source_type"] == "excel_sheet_split_chunk"
    assert chunks[1].order == chunks[0].order + 1


def test_parse_excel_empty_workbook_yields_no_chunks() -> None:
    buf = io.BytesIO()
    openpyxl.Workbook().save(buf)  # a fresh workbook's single default sheet is empty
    assert excel.parse_excel(buf.getvalue()) == []


def test_parse_excel_multiple_sheets_order_by_sheet_then_segment() -> None:
    def build(wb: openpyxl.Workbook) -> None:
        first = wb.active
        first.title = "First"
        first.append(["A", "B"])
        for i in range(2):
            first.append([i, i])
        second = wb.create_sheet("Second")
        second.append(["A", "B"])
        for i in range(2):
            second.append([i, i])

    chunks = excel.parse_excel(_workbook_bytes(build))
    assert len(chunks) == 2
    assert chunks[0].order < chunks[1].order
    assert chunks[0].metadata["sheet_name"] == "First"
    assert chunks[1].metadata["sheet_name"] == "Second"


# --------------------------------------------------------------------------- #
# docx — heading detection, block merging, tables, document order             #
# --------------------------------------------------------------------------- #
def _docx_bytes(build: Callable[[WordDocumentObject], None]) -> bytes:
    document = WordDocument()
    build(document)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _add_table(document: WordDocumentObject, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            table.cell(r, c).text = cell


def test_parse_docx_merges_paragraphs_under_their_heading() -> None:
    """Alpha's merge rule: a heading opens a block and the paragraphs below it
    join that block, so a section arrives as one coherent chunk."""

    def build(document: WordDocumentObject) -> None:
        document.add_heading("Attendance Policy", level=1)
        document.add_paragraph("Employees must record arrival and departure daily.")
        document.add_paragraph("- Late arrivals are logged by the system.")

    chunks = docx_parser.parse_docx(_docx_bytes(build))

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.kind is ParsedChunkKind.TEXT
    assert chunk.metadata["title"] == "Attendance Policy"
    assert chunk.metadata["section_type"] == "section"
    # The heading is document text, not injected metadata: dropping it here
    # would keep it out of the index entirely (unlike a PDF table caption,
    # which the PDF text pass indexes on its own).
    assert chunk.text.startswith("Attendance Policy\n")
    assert "Late arrivals" in chunk.text


def test_parse_docx_detects_a_heading_by_word_style_in_any_language() -> None:
    """Decision س-09 = أ: the style is a structural signal, so an Arabic
    heading is detected exactly as an English one is — which alpha's
    English-keyword heuristic cannot do (fact ح-١٤)."""

    def build(document: WordDocumentObject) -> None:
        document.add_heading("سياسة الحضور والانصراف.", level=2)
        document.add_paragraph("على الموظّفين تسجيل الحضور والانصراف يوميًّا.")

    chunks = docx_parser.parse_docx(_docx_bytes(build))

    assert len(chunks) == 1
    assert chunks[0].metadata["title"] == "سياسة الحضور والانصراف."
    # The text heuristic alone would have refused it twice over: it ends with
    # a full stop, and none of its words are in the English keyword list.
    assert docx_parser._is_title("سياسة الحضور والانصراف.") is False


def test_parse_docx_falls_back_to_the_text_heuristic_without_styles() -> None:
    """An unstyled document still yields sections: the heuristic answers for
    every paragraph that carries no heading style."""

    def build(document: WordDocumentObject) -> None:
        document.add_paragraph("Purpose:")
        document.add_paragraph("To define how attendance is recorded and reviewed.")
        document.add_paragraph("2.1 Scope")
        document.add_paragraph("This section applies to all staff of the company.")

    chunks = docx_parser.parse_docx(_docx_bytes(build))

    assert [c.metadata["title"] for c in chunks] == ["Purpose:", "2.1 Scope"]


@pytest.mark.parametrize(
    ("text", "expected"),
    [
        ("Purpose:", True),  # short label ending in a colon
        ("HR & Admin: the team is responsible for:", False),  # a lead-in, not a heading
        ("2.3.1 Scope", True),  # multi-level numbering
        ("Annex 1", True),  # sectioned heading
        ("responsibilities", True),  # known keyword
        ("Job Analysis", True),  # short capitalized line
        ("GENERAL PROVISIONS OF THE POLICY", True),  # ALL CAPS
        ("Internal and External Communication Guidelines", True),  # Title Case
        ("Employees must record their arrival time daily.", False),  # a sentence
        ("Class A1 | 150000 | A", False),  # tabular data
        ("", False),
    ],
)
def test_is_title_ladder(text: str, expected: bool) -> None:
    """Alpha's `is_title` ladder, rule by rule — the calibration this port
    keeps unchanged."""
    assert docx_parser._is_title(text) is expected


def test_parse_docx_keeps_an_unheaded_paragraph_as_its_own_block() -> None:
    """Body text with no heading above it stands alone rather than joining an
    unrelated neighbour."""

    def build(document: WordDocumentObject) -> None:
        document.add_paragraph("The first standalone sentence of the document.")
        document.add_paragraph("A second, entirely unrelated standalone sentence.")

    chunks = docx_parser.parse_docx(_docx_bytes(build))

    assert len(chunks) == 2
    assert all(c.metadata.get("title") is None for c in chunks)
    assert all(c.metadata["section_type"] == "paragraph" for c in chunks)


def test_parse_docx_stamps_a_block_with_its_own_first_position() -> None:
    """The divergence from alpha: alpha stamps a block with the position of
    the heading that ENDED it, which makes `position_in_doc` non-monotonic —
    and step 10 (`P-17`) consumes it as an ordering signal."""

    def build(document: WordDocumentObject) -> None:
        document.add_heading("First Section", level=1)
        document.add_paragraph("Body of the first section, long enough to matter.")
        document.add_heading("Second Section", level=1)
        document.add_paragraph("Body of the second section, also long enough.")

    chunks = docx_parser.parse_docx(_docx_bytes(build))

    assert [c.metadata["position_in_doc"] for c in chunks] == [0, 2]
    assert [c.metadata["paragraph_number"] for c in chunks] == [0, 2]
    assert [c.order for c in chunks] == sorted(c.order for c in chunks)


def test_parse_docx_splits_a_long_block_and_repeats_the_heading() -> None:
    """A block over `MAX_CHUNK_CHARS` is split, and every part carries the
    heading again: a part retrieved on its own with no section name above it
    is a part no reader can place."""
    body = "Every employee is expected to follow the stated procedure. " * 80

    def build(document: WordDocumentObject) -> None:
        document.add_heading("Procedures", level=1)
        document.add_paragraph(body.strip())

    chunks = docx_parser.parse_docx(_docx_bytes(build))

    assert len(chunks) > 1
    assert all(c.text.startswith("Procedures\n") for c in chunks)
    assert all(len(c.text) <= docx_parser.MAX_CHUNK_CHARS for c in chunks)
    # Parts stay adjacent on the order axis and never reach the next item.
    assert [c.metadata["part_index"] for c in chunks] == list(range(len(chunks)))
    assert [c.order for c in chunks] == [chunks[0].order + i for i in range(len(chunks))]
    assert all(c.metadata["part_count"] == len(chunks) for c in chunks)


def test_parse_docx_extracts_a_table_with_its_heading_breadcrumb() -> None:
    """A table's title is section + caption, so a generic caption never
    shadows the policy the table belongs to."""

    def build(document: WordDocumentObject) -> None:
        document.add_heading("Attendance Policy", level=1)
        document.add_paragraph("Quick Reference Table:")
        _add_table(document, [["Grade", "Allowance"], ["A1", "1500"], ["B2", "900"]])

    chunks = docx_parser.parse_docx(_docx_bytes(build))
    tables = [c for c in chunks if c.kind is ParsedChunkKind.TABLE]

    assert len(tables) == 1
    table = tables[0]
    assert table.metadata["title"] == "Attendance Policy — Quick Reference Table:"
    assert table.metadata["headers"] == ["Grade", "Allowance"]
    assert table.metadata["num_rows"] == 2
    assert table.metadata["num_cols"] == 2
    assert json.loads(table.text)["rows"] == [
        {"Grade": "A1", "Allowance": "1500"},
        {"Grade": "B2", "Allowance": "900"},
    ]


def test_parse_docx_skips_a_table_without_data_rows() -> None:
    """A header row alone is not a table (alpha drops it too)."""

    def build(document: WordDocumentObject) -> None:
        _add_table(document, [["Grade", "Allowance"]])

    assert docx_parser.parse_docx(_docx_bytes(build)) == []


def test_parse_docx_names_an_empty_header_cell_by_position() -> None:
    def build(document: WordDocumentObject) -> None:
        _add_table(document, [["Grade", ""], ["A1", "1500"]])

    table = docx_parser.parse_docx(_docx_bytes(build))[0]

    assert table.metadata["headers"] == ["Grade", "Column_2"]
    assert json.loads(table.text)["rows"] == [{"Grade": "A1", "Column_2": "1500"}]


def test_parse_docx_orders_text_and_tables_in_document_order() -> None:
    """Paragraphs and tables come out of ONE walk of the document, so a table
    sorts between the sections it sits between."""

    def build(document: WordDocumentObject) -> None:
        document.add_heading("First Section", level=1)
        document.add_paragraph("Body of the first section, long enough to matter.")
        _add_table(document, [["Grade", "Allowance"], ["A1", "1500"]])
        document.add_heading("Second Section", level=1)
        document.add_paragraph("Body of the second section, also long enough.")

    chunks = sorted(docx_parser.parse_docx(_docx_bytes(build)), key=lambda c: c.order)

    assert [c.kind for c in chunks] == [
        ParsedChunkKind.TEXT,
        ParsedChunkKind.TABLE,
        ParsedChunkKind.TEXT,
    ]
    assert [c.metadata["position_in_doc"] for c in chunks] == [0, 2, 3]


def test_parse_docx_empty_document_yields_no_chunks() -> None:
    assert docx_parser.parse_docx(_docx_bytes(lambda document: None)) == []


def test_parse_docx_raises_on_bytes_that_are_not_a_word_document() -> None:
    """Unlike a PDF there is nothing to degrade to: a DOCX opens or it does
    not, and the extractor turns the failure into a `ValidationError`."""
    with pytest.raises(Exception):  # noqa: B017 — python-docx raises its own types
        docx_parser.parse_docx(b"not a word document at all")


@pytest.mark.parametrize(
    ("table_count", "text_count", "expected"),
    [
        (0, 0, "docx_empty"),
        (3, 0, "docx_structured_text"),
        (0, 3, "docx_unstructured_text"),
        (7, 2, "docx_structured_text"),
        (2, 5, "docx_semi_structured_text"),
        (1, 9, "docx_unstructured_text"),
    ],
)
def test_classify_docx_thresholds(table_count: int, text_count: int, expected: str) -> None:
    assert docx_parser.classify_docx(table_count=table_count, text_count=text_count) == expected


# --------------------------------------------------------------------------- #
# pdf_text — block granularity, structural order, noise guards                 #
# --------------------------------------------------------------------------- #
# Each `insert_text` at its own y position lays out as its own fitz block, so
# a page's block list is exactly the lines written to it, in that order.
def _blocks_pdf(pages: list[list[str]]) -> bytes:
    doc = fitz.open()
    for lines in pages:
        page = doc.new_page()
        for index, line in enumerate(lines):
            page.insert_text((72, 72 + index * 120), line)
    data: bytes = doc.tobytes()
    doc.close()
    return data


def test_parse_pdf_text_single_block() -> None:
    text = (
        "Hello world, this is a longer test page with enough characters. "
        "عربي نص تجريبي طويل بما فيه الكفاية."
    )
    chunks = pdf_text.parse_pdf_text(_blocks_pdf([[text]]))

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.order == 0
    assert chunk.kind is ParsedChunkKind.TEXT
    assert chunk.metadata["page_number"] == 1
    assert chunk.metadata["position_in_doc"] == 0
    assert chunk.metadata["source_ext"] == ".pdf"
    assert "Hello world" in chunk.text


def test_parse_pdf_text_emits_one_chunk_per_block_not_per_page() -> None:
    """The granularity change of plan step 2 (decision س-08): a page with three
    paragraphs is three chunks, each carrying only its own paragraph."""
    chunks = pdf_text.parse_pdf_text(
        _blocks_pdf(
            [
                [
                    "First paragraph block with plenty of characters here.",
                    "Second paragraph block, also long enough to survive.",
                    "Third paragraph block, still well past the threshold.",
                ]
            ]
        )
    )

    assert len(chunks) == 3
    assert [c.order for c in chunks] == [0, 1, 2]
    assert [c.metadata["position_in_doc"] for c in chunks] == [0, 1, 2]
    assert [c.metadata["chunk_index"] for c in chunks] == [0, 1, 2]
    assert all(c.metadata["page_number"] == 1 for c in chunks)
    assert chunks[0].text.startswith("First")
    assert "Second" not in chunks[0].text


def test_parse_pdf_text_orders_blocks_on_a_page_strided_axis() -> None:
    """`order` is a document-wide block rank: page 2's first block sorts after
    every block of page 1, on the same axis `pdf_tables.py` emits on."""
    chunks = pdf_text.parse_pdf_text(
        _blocks_pdf(
            [
                [
                    "Page one, block one, comfortably above the threshold.",
                    "Page one, block two, also above the threshold.",
                ],
                ["Page two, block one, comfortably above the threshold."],
            ]
        )
    )

    assert [c.order for c in chunks] == [0, 1, 1000]
    assert [c.metadata["page_number"] for c in chunks] == [1, 1, 2]
    assert [c.metadata["block_index"] for c in chunks] == [0, 1, 0]
    assert [c.order for c in chunks] == sorted(c.order for c in chunks)


def test_parse_pdf_text_skips_blocks_below_min_block_chars() -> None:
    """The threshold moved from the page to the block, and dropping a block
    must NOT renumber the ones that survive it — `order` and `position_in_doc`
    are structural, so the surviving block keeps the rank of its position."""
    chunks = pdf_text.parse_pdf_text(
        _blocks_pdf([["short", "A second block that is long enough to be kept."]])
    )

    assert len(chunks) == 1
    assert chunks[0].order == 1
    assert chunks[0].metadata["block_index"] == 1
    assert chunks[0].metadata["position_in_doc"] == 1
    assert chunks[0].metadata["chunk_index"] == 0


def test_parse_pdf_text_skips_a_page_of_only_short_blocks() -> None:
    assert pdf_text.parse_pdf_text(_blocks_pdf([["short", "tiny"]])) == []


def test_parse_pdf_text_carries_the_block_rectangle_in_fitz_coordinates() -> None:
    chunks = pdf_text.parse_pdf_text(
        _blocks_pdf([["A block whose rectangle should be reported back."]])
    )

    bbox = chunks[0].metadata["block_bbox"]
    assert len(bbox) == 4
    assert all(isinstance(v, float) for v in bbox)
    # fitz origin is top-left with y growing downward, so y0 < y1 and a block
    # written near the top of the page has a small y0.
    assert bbox[0] < bbox[2]
    assert bbox[1] < bbox[3]
    assert bbox[1] < 100


@pytest.mark.parametrize(
    "block",
    [
        pytest.param((0.0, 0.0, 1.0, 1.0), id="malformed_tuple_too_short"),
        pytest.param(
            (0.0, 0.0, 1.0, 1.0, "<image: DeviceRGB, width 800, height 600, bpc 8>", 0, 1),
            id="image_block",
        ),
    ],
)
def test_block_text_rejects_unusable_blocks(block: tuple[object, ...]) -> None:
    """Both guards are defensive — the pinned PyMuPDF emits neither shape — but
    an image block's synthetic placeholder is longer than MIN_BLOCK_CHARS and
    would be indexed as prose if the `blocks` flags ever changed."""
    assert pdf_text._block_text(block) == ""


def _region(page_number: int, rect: fitz.Rect, page_height: float) -> pdf_tables.TableRegion:
    """A `TableRegion` covering `rect` (given in fitz coordinates), expressed
    the way camelot reports one: PDF coordinates, origin bottom-left."""
    return pdf_tables.TableRegion(
        page_number=page_number,
        bbox=(rect.x0, page_height - rect.y1, rect.x1, page_height - rect.y0),
    )


def _page_height(data: bytes) -> float:
    doc = fitz.open(stream=data, filetype="pdf")
    height = float(doc[0].rect.height)
    doc.close()
    return height


def _block_rect(data: bytes, block_index: int) -> fitz.Rect:
    doc = fitz.open(stream=data, filetype="pdf")
    block = doc[0].get_text("blocks", sort=True)[block_index]
    doc.close()
    return fitz.Rect(block[:4])


def test_parse_pdf_text_drops_a_block_inside_a_table_region() -> None:
    """Plan step 3 (`P-07`): the table pass already indexed that region as a
    structured table, so re-indexing it here as flattened prose would put the
    same content in the index twice."""
    data = _blocks_pdf(
        [
            [
                "A paragraph of ordinary prose, well above the threshold.",
                "Name Dept Salary rows that camelot boxed as a table.",
            ]
        ]
    )
    regions = {1: [_region(1, _block_rect(data, 1), _page_height(data))]}

    chunks = pdf_text.parse_pdf_text(data, table_regions=regions)

    assert len(chunks) == 1
    assert chunks[0].text.startswith("A paragraph")
    assert chunks[0].metadata["tables_avoided"] == 1
    # Dropping a block never renumbers what survives it: the kept block is
    # still block 0 at position 0, exactly as with no regions at all.
    assert chunks[0].order == 0
    assert chunks[0].metadata["position_in_doc"] == 0


def test_parse_pdf_text_keeps_a_block_that_only_grazes_a_table_region() -> None:
    """The test is on overlapping AREA, not on contact: camelot's bbox is
    generous, and a paragraph brushing a table's edge is still a paragraph."""
    data = _blocks_pdf([["A paragraph that merely touches the table's edge and stays prose."]])
    rect = _block_rect(data, 0)
    grazing = fitz.Rect(rect.x0, rect.y1 - rect.height * 0.2, rect.x1, rect.y1 + 40)
    regions = {1: [_region(1, grazing, _page_height(data))]}

    chunks = pdf_text.parse_pdf_text(data, table_regions=regions)

    assert len(chunks) == 1
    assert chunks[0].metadata["tables_avoided"] == 1


def test_parse_pdf_text_drops_a_table_only_page_entirely() -> None:
    """Alpha's "table-only page" rule: what a table's bbox failed to cover is
    stray cell text, not prose — below `TABLE_PAGE_MIN_TEXT_CHARS` in total it
    all goes."""
    data = _blocks_pdf(
        [
            [
                "Name Dept Salary rows that camelot boxed as a table.",
                "Total: 18300 riyals.",
            ]
        ]
    )
    regions = {1: [_region(1, _block_rect(data, 0), _page_height(data))]}

    assert pdf_text.parse_pdf_text(data, table_regions=regions) == []


def test_parse_pdf_text_keeps_a_sparse_page_that_had_no_tables() -> None:
    """The table-only rule fires only where tables were avoided — an ordinary
    page with one short-ish paragraph keeps it."""
    sparse = "One paragraph, just past the block threshold."
    assert len(sparse) < pdf_text.TABLE_PAGE_MIN_TEXT_CHARS

    chunks = pdf_text.parse_pdf_text(_blocks_pdf([[sparse]]))

    assert len(chunks) == 1
    assert chunks[0].metadata["tables_avoided"] == 0


def test_parse_pdf_text_ignores_regions_on_other_pages() -> None:
    data = _blocks_pdf(
        [
            ["Page one block, comfortably above the threshold, kept as is."],
            ["Page two block, comfortably above the threshold, kept as is."],
        ]
    )
    # The region is keyed to page 2 and sits low on it, far from either page's
    # block: page 1 must be untouched, and page 2 must keep the block that does
    # not overlap it.
    height = _page_height(data)
    regions = {2: [_region(2, fitz.Rect(72, height - 200, 400, height - 80), height)]}

    chunks = pdf_text.parse_pdf_text(data, table_regions=regions)

    assert [c.metadata["page_number"] for c in chunks] == [1, 2]
    assert [c.metadata["tables_avoided"] for c in chunks] == [0, 1]


def test_parse_pdf_text_skips_an_encrypted_pdf() -> None:
    doc = fitz.open()
    doc.new_page().insert_text((72, 72), "A secret block, long enough to be kept.")
    data = doc.tobytes(encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw="owner", user_pw="user")
    doc.close()

    assert pdf_text.parse_pdf_text(data) == []


def test_parse_pdf_text_degrades_gracefully_on_corrupt_bytes() -> None:
    assert pdf_text.parse_pdf_text(b"not a pdf") == []


# --------------------------------------------------------------------------- #
# pdf_tables — stream detection, noise guards, cross-page merge, captions      #
# --------------------------------------------------------------------------- #
# Camelot's `stream` flavor infers columns from whitespace, so a fixture only
# has to lay text out in aligned columns. `caption_gap` controls how far above
# the table the caption sits: too close and `edge_tol=50` swallows it INTO the
# table (its text becomes row 0, i.e. the headers) — a real property of the
# flavor, not of this port, and the reason the merge fixture below has no
# caption at all.
_TABLE_COLUMN_X = (72, 202, 332)


def _table_page(doc: fitz.Document, rows: list[tuple[str, ...]], caption: str | None) -> None:
    page = doc.new_page()
    if caption is not None:
        page.insert_text((72, 100), caption, fontsize=13)
    y = 200.0
    for row in rows:
        for x, cell in zip(_TABLE_COLUMN_X, row, strict=False):
            page.insert_text((x, y), cell, fontsize=11)
        y += 22


def _table_pdf_bytes(pages: list[list[tuple[str, ...]]], caption: str | None = None) -> bytes:
    doc = fitz.open()
    for rows in pages:
        _table_page(doc, rows, caption)
    data: bytes = doc.tobytes()
    doc.close()
    return data


_SALARY_ROWS = [
    ("Name", "Dept", "Salary"),
    ("Ali", "Finance", "5000"),
    ("Sara", "HR", "6200"),
    ("Omar", "IT", "7100"),
]


def test_parse_pdf_tables_extracts_a_structured_table() -> None:
    chunks, regions = pdf_tables.parse_pdf_tables(_table_pdf_bytes([_SALARY_ROWS]))

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.kind is ParsedChunkKind.TABLE
    assert chunk.order == 0
    assert chunk.metadata["page_number"] == 1
    assert chunk.metadata["source_ext"] == ".pdf"
    assert chunk.metadata["section_type"] == "structured_table"
    assert chunk.metadata["layout_mode"] == "camelot_stream"
    assert chunk.metadata["is_merged_table"] is False
    assert chunk.metadata["headers"] == ["Name", "Dept", "Salary"]
    assert chunk.metadata["total_rows"] == 3

    payload = json.loads(chunk.text)
    assert payload["headers"] == ["Name", "Dept", "Salary"]
    assert payload["rows"][0] == {"Name": "Ali", "Dept": "Finance", "Salary": "5000"}

    assert list(regions) == [1]
    assert regions[1][0].page_number == 1
    assert len(regions[1][0].bbox) == 4


def test_parse_pdf_tables_reads_the_caption_above_the_table_into_title() -> None:
    data = _table_pdf_bytes([_SALARY_ROWS], caption="Employee Salary Table")

    chunks, _regions = pdf_tables.parse_pdf_tables(data)

    assert len(chunks) == 1
    # The caption is metadata only — it is deliberately NOT prepended to the
    # chunk text (plan §7: no metadata injection into node text).
    assert chunks[0].metadata["title"] == "Employee Salary Table"
    assert "Employee Salary Table" not in chunks[0].text


def test_parse_pdf_tables_merges_a_table_spanning_two_pages() -> None:
    data = _table_pdf_bytes(
        [
            [("Name", "Dept", "Salary"), ("Ali", "Finance", "5000"), ("Sara", "HR", "6200")],
            [("Name", "Dept", "Salary"), ("Omar", "IT", "7100"), ("Lina", "Legal", "6800")],
        ]
    )

    chunks, regions = pdf_tables.parse_pdf_tables(data)

    # One table, not two — the whole point of the cross-page merge.
    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.metadata["is_merged_table"] is True
    assert chunk.metadata["source_pages"] == [1, 2]
    assert chunk.metadata["total_source_tables"] == 2
    assert chunk.metadata["headers"] == ["Name", "Dept", "Salary"]

    # The repeated header on page 2 is dropped, not merged in as a data row.
    names = [row["Name"] for row in json.loads(chunk.text)["rows"]]
    assert names == ["Ali", "Sara", "Omar", "Lina"]

    # A merged table occupies a region on every page it spans — the text pass
    # must avoid all of them (plan step 3).
    assert sorted(regions) == [1, 2]


def test_parse_pdf_tables_drops_prose_misdetected_as_a_table() -> None:
    doc = fitz.open()
    page = doc.new_page()
    y = 100.0
    for _ in range(12):
        page.insert_text(
            (72, y),
            "This is an ordinary paragraph of running prose with no table structure.",
            fontsize=11,
        )
        y += 20
    data = doc.tobytes()
    doc.close()

    # A single column of prose must not survive as a table: it would be indexed
    # garbled AND removed from the text pass, which avoids table regions.
    assert pdf_tables.parse_pdf_tables(data) == ([], {})


# --- the false-table guard (decision س-30) --------------------------------- #
def _toc_pdf_bytes(entries: list[tuple[str, str]]) -> bytes:
    """A table of contents laid out as camelot sees one: a title column ruled
    with dot leaders, and a page-number column far enough right that `stream`
    reads two columns rather than one."""
    doc = fitz.open()
    page = doc.new_page()
    y = 120.0
    for title, page_no in entries:
        page.insert_text((72, y), title, fontsize=11)
        page.insert_text((430, y), page_no, fontsize=11)
        y += 22
    data: bytes = doc.tobytes()
    doc.close()
    return data


_TOC_ENTRIES = [
    ("Table of Contents", "Page"),
    ("1. Introduction ...........", "5"),
    ("2. Scope and Purpose ......", "8"),
    ("3. Mine Auxiliary Trafo ...", "14"),
    ("4. Protection Settings ....", "21"),
    ("5. Commissioning ..........", "30"),
    ("6. Appendix A .............", "44"),
]


def test_parse_pdf_tables_drops_a_table_of_contents_but_keeps_its_region() -> None:
    """س-30. A ToC's rows read exactly like the questions users ask and answer
    them with a page number, so they win the sparse leg where it is asked and
    return nothing — in `k` slots that are frozen and cannot be bought back.

    The region survives the drop, which is what separates this guard from the
    two in `_filter_noise`: those hand the page back to the text pass because
    real prose is underneath, whereas re-indexing a ToC as paragraphs would
    carry the same lexical harm across merely reshaped.
    """
    chunks, regions = pdf_tables.parse_pdf_tables(_toc_pdf_bytes(_TOC_ENTRIES))

    assert chunks == []
    assert list(regions) == [1]
    assert regions[1][0].page_number == 1


def test_a_table_of_contents_survives_neither_pass() -> None:
    """Where the guard's value actually is — not in the table pass alone.

    Had the ToC's region been dropped along with its chunk, as the guards in
    `_filter_noise` drop theirs, `pdf_text` would hand the same lines straight
    back as prose: the lexical harm intact, merely reshaped. Keeping the region
    lets the "table-only page" rule finish the job.
    """
    data = _toc_pdf_bytes(_TOC_ENTRIES)

    table_chunks, regions = pdf_tables.parse_pdf_tables(data)

    assert table_chunks == []
    assert pdf_text.parse_pdf_text(data, table_regions=regions) == []

    # The counterfactual, so this test cannot pass for want of any text at all.
    leaked = pdf_text.parse_pdf_text(data)
    assert leaked
    assert any("Mine Auxiliary Trafo" in chunk.text for chunk in leaked)


def test_parse_pdf_tables_keeps_a_table_whose_dot_rows_stay_under_the_ratio() -> None:
    """A dot run is not by itself disqualifying — a placeholder cell in real
    data must not cost the table."""
    data = _table_pdf_bytes(
        [
            [
                ("Item", "Notes", "Price"),
                ("Cable", "..........", "120"),
                ("Switch", "spare", "340"),
                ("Panel", "spare", "980"),
                ("Relay", "spare", "60"),
            ]
        ]
    )

    chunks, _regions = pdf_tables.parse_pdf_tables(data)

    assert len(chunks) == 1
    assert chunks[0].metadata["total_rows"] == 4


@pytest.mark.parametrize(
    ("rows", "expected"),
    [
        # A full ToC.
        ([["1. Intro ........", "5"], ["2. Scope ........", "8"]], True),
        # Exactly at the ratio: 0.5 drops, the comparison is inclusive.
        ([["1. Intro ........", "5"], ["Cable", "120"]], True),
        # One dot row in three is a placeholder, not a ToC.
        ([["1. Intro ........", "5"], ["Cable", "120"], ["Panel", "980"]], False),
        # Below six dots is an ellipsis or a decimal chain, not a leader.
        ([["Ellipsis ...", "5"], ["Rev 1.2.3.4.5.6", "8"]], False),
        # `stream` cut the leader at the column edge; joining restores it.
        ([["1. Intro ...", "... 5"], ["2. Scope ...", "... 8"]], True),
    ],
)
def test_is_table_of_contents_ratio(rows: list[list[str]], expected: bool) -> None:
    df = pd.DataFrame(rows, columns=["title", "page"])
    assert pdf_tables._is_table_of_contents(df) is expected


def test_is_table_of_contents_is_false_for_an_empty_frame() -> None:
    """No rows is no evidence — and `0/0` must not raise on the way to saying
    so."""
    assert pdf_tables._is_table_of_contents(pd.DataFrame()) is False


def test_parse_pdf_tables_skips_an_encrypted_pdf() -> None:
    doc = fitz.open()
    _table_page(doc, _SALARY_ROWS, None)
    data: bytes = doc.tobytes(encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="secret")
    doc.close()

    assert pdf_tables.parse_pdf_tables(data) == ([], {})


def test_parse_pdf_tables_degrades_gracefully_on_corrupt_bytes() -> None:
    assert pdf_tables.parse_pdf_tables(b"not a pdf") == ([], {})


@pytest.mark.parametrize(
    ("cols1", "cols2", "expected"),
    [
        (["Name", "Dept"], ["Name", "Dept"], 1.0),  # identical
        (["Name", "Dept"], ["name", " dept "], 1.0),  # case/whitespace insensitive
        (["Name", "Dept"], ["Employee Name", "Dept"], 0.85),  # containment scores 0.7
        (["Name", "Dept"], ["Salary", "Year"], 0.0),  # unrelated
        (["Name", "Dept"], ["Name"], 0.0),  # different arity is never a match
        ([], [], 0.0),
    ],
)
def test_column_similarity_calibration(cols1: list[str], cols2: list[str], expected: float) -> None:
    assert pdf_tables._column_similarity(cols1, cols2) == pytest.approx(expected)


# --------------------------------------------------------------------------- #
# image_ocr — dedup, quality gate, upscaling, graceful degradation            #
# --------------------------------------------------------------------------- #
def test_sha1_digest_is_stable_and_distinguishes_content() -> None:
    a = image_ocr.sha1_digest(b"same-bytes")
    b = image_ocr.sha1_digest(b"same-bytes")
    c = image_ocr.sha1_digest(b"different-bytes")
    assert a == b
    assert a != c


@pytest.mark.parametrize(
    ("text", "expected"),
    [
        ("هذا نص عربي حقيقي وواضح", True),
        ("!@#$ %^&* () __ --", False),
        ("hi", False),
        ("", False),
    ],
)
def test_is_meaningful_text(text: str, expected: bool) -> None:
    assert image_ocr.is_meaningful_text(text) is expected


def test_upscale_if_small_doubles_small_images() -> None:
    img = Image.new("L", (100, 50), color=255)
    assert image_ocr._upscale_if_small(img).size == (200, 100)


def test_upscale_if_small_leaves_large_images_untouched() -> None:
    img = Image.new("L", (1200, 800), color=255)
    assert image_ocr._upscale_if_small(img).size == (1200, 800)


def test_run_ocr_upscales_before_calling_tesseract(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, object] = {}

    def fake_image_to_string(
        img: Image.Image, lang: str | None = None, config: str | None = None
    ) -> str:
        captured["size"] = img.size
        captured["lang"] = lang
        return "  stubbed text  "

    monkeypatch.setattr(image_ocr.pytesseract, "image_to_string", fake_image_to_string)

    result = image_ocr.run_ocr(Image.new("L", (100, 50), color=255))

    assert result == "stubbed text"
    assert captured["size"] == (200, 100)
    assert captured["lang"] == image_ocr.OCR_LANG


def test_run_ocr_returns_empty_string_when_tesseract_missing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def raise_not_found(
        img: Image.Image, lang: str | None = None, config: str | None = None
    ) -> str:
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(image_ocr.pytesseract, "image_to_string", raise_not_found)

    assert image_ocr.run_ocr(Image.new("L", (50, 50), color=255)) == ""


def test_run_ocr_returns_empty_string_on_generic_ocr_error(monkeypatch: pytest.MonkeyPatch) -> None:
    def raise_generic(img: Image.Image, lang: str | None = None, config: str | None = None) -> str:
        raise RuntimeError("boom")

    monkeypatch.setattr(image_ocr.pytesseract, "image_to_string", raise_generic)

    assert image_ocr.run_ocr(Image.new("L", (50, 50), color=255)) == ""


def test_parse_image_blank_image_yields_no_chunk() -> None:
    if shutil.which("tesseract") is None:
        pytest.skip("tesseract binary not available")

    img = Image.new("RGB", (200, 200), color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    assert image_ocr.parse_image(buf.getvalue()) == []


def test_parse_image_degrades_gracefully_on_corrupt_bytes() -> None:
    assert image_ocr.parse_image(b"not an image") == []


# --------------------------------------------------------------------------- #
# image_ocr — embedded images: grouping, dedup, caps, placeholders (step 5)   #
# --------------------------------------------------------------------------- #
_OCR_TEXT = "TOTAL REVENUE FOR 2026 IS 1,240,000 RIYAL"


def _png(width: int, height: int, *, tint: int = 255) -> bytes:
    """A distinct-per-`tint` PNG: two calls with different tints are different
    bytes, which is what the SHA-1 dedup tests need to control."""
    buf = io.BytesIO()
    Image.new("RGB", (width, height), color=(tint, tint, tint)).save(buf, format="PNG")
    return buf.getvalue()


def _stub_ocr(
    monkeypatch: pytest.MonkeyPatch, *, readable: Callable[[Image.Image], bool] | None = None
) -> None:
    """Replace the tesseract call. The binary may be installed without the
    Arabic language data, so no test here may depend on what it would return —
    only on what this module does with the result."""
    decides = readable if readable is not None else (lambda img: True)
    monkeypatch.setattr(image_ocr, "run_ocr", lambda img: _OCR_TEXT if decides(img) else "")


def _image_pdf(pages: list[list[bytes]]) -> bytes:
    doc = fitz.open()
    for images in pages:
        page = doc.new_page()
        top = 50.0
        for blob in images:
            page.insert_image(fitz.Rect(50, top, 250, top + 120), stream=blob)
            top += 130
    data: bytes = doc.tobytes()
    doc.close()
    return data


def _image_docx(images: list[bytes]) -> bytes:
    return _docx_bytes(
        lambda document: [document.add_picture(io.BytesIO(blob)) for blob in images] and None
    )


def _image_xlsx(images: list[bytes]) -> bytes:
    def build(workbook: openpyxl.Workbook) -> None:
        sheet = workbook.active
        assert sheet is not None
        for index, blob in enumerate(images):
            sheet.add_image(WorksheetImage(io.BytesIO(blob)), f"B{2 + index * 20}")

    return _workbook_bytes(build)


def test_clean_ocr_text_collapses_the_noise_a_scanner_adds() -> None:
    raw = "Title\n\n\n\nBody   with    gaps\n______________\nEnd  "

    assert image_ocr.clean_ocr_text(raw) == "Title\n\nBody with gaps\n\nEnd"


def test_clean_ocr_text_of_nothing_is_nothing() -> None:
    assert image_ocr.clean_ocr_text("") == ""


def test_create_page_summary_counts_both_kinds_of_image() -> None:
    images = [{"has_meaningful_text": True}, {"has_meaningful_text": False}]

    summary = image_ocr.create_page_summary(images)

    assert summary == (
        "This page contains 1 image(s) with extracted text. 1 diagram(s)/chart(s) without text."
    )


def test_create_page_summary_is_empty_without_images() -> None:
    assert image_ocr.create_page_summary([]) == ""


def test_pdf_images_merge_into_one_chunk_per_page(monkeypatch: pytest.MonkeyPatch) -> None:
    """A page's images are one figure more often than they are several, so the
    group — not the image — is the chunk. The order is the page-strided axis
    the other two PDF passes emit on, with 999 as the page's last slot."""
    _stub_ocr(monkeypatch)
    data = _image_pdf(
        [[_png(400, 400, tint=10), _png(400, 400, tint=20)], [_png(400, 400, tint=30)]]
    )

    result = image_ocr.parse_pdf_images(data)

    assert result.image_count == 3
    assert result.truncated is False
    assert [chunk.order for chunk in result.chunks] == [999, 1999]
    assert [chunk.kind for chunk in result.chunks] == [ParsedChunkKind.OCR] * 2
    first = result.chunks[0]
    assert first.metadata["page_number"] == 1
    assert first.metadata["image_count"] == 2
    assert first.metadata["section_type"] == "page_merged_images"
    assert first.text == f"{_OCR_TEXT}\n\n{_OCR_TEXT}"


def test_an_image_below_the_size_filter_is_never_ocred(monkeypatch: pytest.MonkeyPatch) -> None:
    """`ocr_min_image_px` is the icon/logo guard (§3.8) — and it runs before
    OCR, so a rejected image costs a header read, not a tesseract call."""
    calls: list[int] = []
    monkeypatch.setattr(image_ocr, "run_ocr", lambda img: calls.append(img.width) or _OCR_TEXT)
    data = _image_pdf([[_png(150, 150)]])

    result = image_ocr.parse_pdf_images(data)

    assert result.chunks == ()
    assert result.image_count == 0
    assert calls == []


def test_a_logo_repeated_on_every_page_is_read_once(monkeypatch: pytest.MonkeyPatch) -> None:
    """Deduplication is document-wide, not per page — alpha's `seen_hashes`."""
    _stub_ocr(monkeypatch)
    logo = _png(400, 400, tint=77)

    result = image_ocr.parse_pdf_images(_image_pdf([[logo], [logo], [logo]]))

    assert result.image_count == 1
    assert len(result.chunks) == 1
    assert result.chunks[0].metadata["page_number"] == 1


def test_sha1_dedup_spans_groups(monkeypatch: pytest.MonkeyPatch) -> None:
    """The same bytes arriving in a LATER group are dropped there too, which is
    what makes the guarantee above document-wide rather than page-local."""
    _stub_ocr(monkeypatch)
    blob = _png(400, 400)
    groups = [
        image_ocr._Group(
            order=index,
            metadata={"page_number": index + 1},
            images=iter([image_ocr._RawImage(data=blob, metadata={})]),
        )
        for index in range(2)
    ]

    result = image_ocr._ocr_groups(groups, caps=image_ocr._caps(None))

    assert result.image_count == 1
    assert [chunk.metadata["page_number"] for chunk in result.chunks] == [1]


def test_a_textless_image_contributes_its_placeholder_and_a_page_context(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """`P-11`'s placeholder is the LIVE branch here (alpha's is dead code under
    its own default): the merged text says where the figure sits, and the
    `[Page Context: ...]` prefix says what the page turned out to be."""
    _stub_ocr(monkeypatch, readable=lambda img: img.width >= 600)
    data = _image_pdf([[_png(700, 400, tint=10), _png(300, 300, tint=20)]])

    chunk = image_ocr.parse_pdf_images(data).chunks[0]

    assert chunk.text == (
        "[Page Context: This page contains 1 image(s) with extracted text. "
        "1 diagram(s)/chart(s) without text.]\n\n"
        f"{_OCR_TEXT}\n\n[Diagram/Chart 300x300px]"
    )
    assert chunk.metadata["images_with_text"] == 1
    assert chunk.metadata["images_without_text"] == 1


def test_a_page_whose_images_all_read_gets_no_page_context_prefix(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Alpha adds the prefix only when something was lost on THIS page; a clean
    page is never padded with boilerplate."""
    _stub_ocr(monkeypatch)

    chunk = image_ocr.parse_pdf_images(_image_pdf([[_png(400, 400)]])).chunks[0]

    assert chunk.text == _OCR_TEXT
    assert chunk.metadata["images_without_text"] == 0


def test_the_per_page_cap_stops_one_page_eating_the_budget(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _stub_ocr(monkeypatch)
    data = _image_pdf([[_png(400, 400, tint=t) for t in (10, 20, 30)], [_png(400, 400, tint=40)]])

    result = image_ocr.parse_pdf_images(data, limits=Limits(ocr_max_images_per_page=1))

    assert result.image_count == 2  # one per page, not three on the first
    assert result.truncated is True
    assert [chunk.metadata["page_number"] for chunk in result.chunks] == [1, 2]


def test_the_per_document_cap_truncates_and_says_so(monkeypatch: pytest.MonkeyPatch) -> None:
    """A cap is a cost decision, not a defect (§3.8): the pass stops and
    DECLARES it rather than failing the document."""
    _stub_ocr(monkeypatch)
    data = _image_pdf([[_png(400, 400, tint=10)], [_png(400, 400, tint=20)]])

    result = image_ocr.parse_pdf_images(data, limits=Limits(ocr_max_images_per_document=1))

    assert result.image_count == 1
    assert result.truncated is True
    assert len(result.chunks) == 1


def test_docx_media_becomes_one_chunk_after_every_positioned_chunk(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """`word/media/` says nothing about where a picture sits, so the archive's
    whole media set is one group at a sort-last order."""
    _stub_ocr(monkeypatch)

    result = image_ocr.parse_office_images(
        _image_docx([_png(400, 400, tint=10), _png(400, 400, tint=20)]),
        media=image_ocr.OfficeMedia.DOCX,
    )

    assert len(result.chunks) == 1
    chunk = result.chunks[0]
    assert chunk.order == image_ocr._MEDIA_ORDER
    assert chunk.kind is ParsedChunkKind.OCR
    assert chunk.metadata["section_type"] == "embedded_images"
    assert chunk.metadata["chunk_type"] == "docx_media_images"
    assert chunk.metadata["image_count"] == 2
    assert all(
        record["media_name"].startswith("word/media/") for record in chunk.metadata["images"]
    )


def test_xlsx_media_is_read_from_its_own_folder(monkeypatch: pytest.MonkeyPatch) -> None:
    _stub_ocr(monkeypatch)

    result = image_ocr.parse_office_images(
        _image_xlsx([_png(400, 400)]), media=image_ocr.OfficeMedia.XLSX
    )

    chunk = result.chunks[0]
    assert chunk.metadata["chunk_type"] == "xlsx_media_images"
    assert chunk.metadata["images"][0]["media_name"].startswith("xl/media/")


def test_a_docx_without_pictures_yields_nothing(monkeypatch: pytest.MonkeyPatch) -> None:
    _stub_ocr(monkeypatch)
    data = _docx_bytes(lambda document: document.add_paragraph("prose only"))

    assert image_ocr.parse_office_images(data, media=image_ocr.OfficeMedia.DOCX).chunks == ()


def test_office_images_degrades_on_bytes_that_are_not_an_archive() -> None:
    result = image_ocr.parse_office_images(b"not a zip", media=image_ocr.OfficeMedia.DOCX)

    assert result == image_ocr.OcrResult(chunks=(), image_count=0, truncated=False)


def test_pdf_images_degrades_on_corrupt_bytes() -> None:
    """Losing a document's figures must never cost it the prose the other two
    passes already produced."""
    assert image_ocr.parse_pdf_images(b"%PDF-1.4 broken").chunks == ()


def test_parse_image_applies_neither_the_size_filter_nor_the_placeholder(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A standalone image IS the document (divergence 2): a small upload is a
    deliberate act, and a placeholder as the whole document text is noise."""
    _stub_ocr(monkeypatch)
    small = _png(150, 150)

    assert image_ocr.parse_image(small)[0].text == _OCR_TEXT

    monkeypatch.setattr(image_ocr, "run_ocr", lambda img: "")
    assert image_ocr.parse_image(small) == []


def test_the_ocr_caps_come_from_settings_with_the_values_the_plan_chose() -> None:
    """§3.8's three numbers, chosen and written down — alpha reads eleven such
    knobs from the environment and documents none of them."""
    limits = Limits()

    assert (limits.ocr_min_image_px, limits.ocr_max_images_per_document) == (200, 40)
    assert limits.ocr_max_images_per_page == 8
    assert image_ocr._caps(None) == image_ocr._caps(limits)


# --------------------------------------------------------------------------- #
# extractor — dispatch, unsupported type, empty guard, end-to-end assembly    #
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    "filename",
    [
        "report.pdf",
        "sheet.xlsx",
        "data.json",
        "policy.docx",
        "notes.txt",
        "readme.md",
        "table.csv",
        "photo.png",
        "photo.jpg",
        "photo.jpeg",
        "photo.webp",
    ],
)
def test_extractor_supports_every_routed_extension(filename: str) -> None:
    extractor = DocumentContentExtractor()
    assert extractor.supports(content_type="application/octet-stream", filename=filename) is True


@pytest.mark.parametrize("filename", ["archive.xls", "image.gif", "data.bin"])
def test_extractor_does_not_support_deferred_or_unknown_extensions(filename: str) -> None:
    extractor = DocumentContentExtractor()
    assert extractor.supports(content_type="application/octet-stream", filename=filename) is False


def test_extractor_raises_unsupported_type_for_unroutable_extension() -> None:
    extractor = DocumentContentExtractor()
    with pytest.raises(UnsupportedTypeError):
        extractor.extract(
            data=b"whatever", filename="archive.xls", content_type="application/octet-stream"
        )


def test_extractor_raises_validation_error_on_empty_bytes() -> None:
    extractor = DocumentContentExtractor()
    with pytest.raises(ValidationError):
        extractor.extract(data=b"", filename="notes.txt", content_type="text/plain")


def test_extractor_raises_validation_error_on_genuine_parse_failure() -> None:
    extractor = DocumentContentExtractor()
    with pytest.raises(ValidationError):
        extractor.extract(
            data=b"not a real workbook", filename="broken.xlsx", content_type="application/xlsx"
        )


# --------------------------------------------------------------------------- #
# extractor — the zip-bomb guard (§3.7, decision س-13, plan step 14)          #
# --------------------------------------------------------------------------- #
def _zip_bytes(entries: dict[str, bytes], *, compression: int = zipfile.ZIP_DEFLATED) -> bytes:
    """A real, valid zip archive with the given ``{name: content}`` members —
    the "an ordinary Office file" half of the guard's test matrix."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=compression) as archive:
        for name, content in entries.items():
            archive.writestr(name, content)
    return buffer.getvalue()


def test_the_zip_bomb_guard_settings_are_the_values_the_plan_chose() -> None:
    """§3.7's three numbers, chosen and written down — the ``ocr_*`` precedent
    (`test_the_ocr_caps_come_from_settings_with_the_values_the_plan_chose`),
    applied to the guard that runs before any parser touches the bytes."""
    limits = Limits()

    assert limits.parser_max_uncompressed_mb == 512
    assert limits.parser_max_compression_ratio == 100
    assert limits.parser_timeout_seconds == 300


@pytest.mark.parametrize("ext", [".docx", ".xlsx"])
def test_zip_bomb_guard_rejects_a_non_zip_masquerading_as_office(ext: str) -> None:
    """``BadZipFile`` — a file that lies about being an Office document under
    its extension never reaches ``python-docx``/``pandas`` at all."""
    with pytest.raises(ValidationError) as exc_info:
        extractor_module._guard_zip_bomb(b"not a zip archive at all", ext, Limits())

    assert exc_info.value.code == "knowledge.zip_bomb"


def test_zip_bomb_guard_rejects_uncompressed_size_over_the_cap() -> None:
    data = _zip_bytes({"word/document.xml": b"<w/>" * 1_000})
    tiny_cap = Limits(parser_max_uncompressed_mb=0)

    with pytest.raises(ValidationError) as exc_info:
        extractor_module._guard_zip_bomb(data, ".docx", tiny_cap)

    assert exc_info.value.code == "knowledge.zip_bomb"
    assert "uncompressed size" in str(exc_info.value)


def test_zip_bomb_guard_rejects_compression_ratio_over_the_cap() -> None:
    """A legitimate Office file runs 3:1 to 10:1 (§3.7's own table); a
    million bytes of one repeated character compresses far past the 100:1
    bomb threshold while staying well under the (default) size cap."""
    data = _zip_bytes({"word/document.xml": b"A" * 1_000_000})

    with pytest.raises(ValidationError) as exc_info:
        extractor_module._guard_zip_bomb(data, ".docx", Limits())

    assert exc_info.value.code == "knowledge.zip_bomb"
    assert "compression ratio" in str(exc_info.value)


def test_zip_bomb_guard_accepts_an_ordinary_office_archive_within_limits() -> None:
    """The sanity check every rejection test above needs: a real, modestly
    sized member with ordinary text content trips neither guard."""
    data = _zip_bytes({"word/document.xml": b"<w:p>Hello, world.</w:p>" * 5})

    extractor_module._guard_zip_bomb(data, ".docx", Limits())  # does not raise


@pytest.mark.parametrize("ext", [".pdf", ".json", ".txt", ".png"])
def test_zip_bomb_guard_is_a_no_op_for_non_office_extensions(ext: str) -> None:
    """Only ``.docx``/``.xlsx`` are zip archives under this router; every
    other route is never asked to open one, even under an absurd cap."""
    extractor_module._guard_zip_bomb(b"anything at all", ext, Limits(parser_max_uncompressed_mb=0))


def test_extractor_end_to_end_fails_a_document_whose_docx_trips_the_zip_bomb_guard() -> None:
    """The guard fires from ``extract()`` itself, before the docx route runs
    — the same ``ValidationError`` shape a genuine parse failure already
    raises, so no new handling is needed downstream."""
    data = _zip_bytes({"word/document.xml": b"A" * 1_000_000})
    tiny = Limits(parser_max_compression_ratio=10)

    with pytest.raises(ValidationError) as exc_info:
        DocumentContentExtractor(limits=tiny).extract(
            data=data,
            filename="bomb.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    assert exc_info.value.code == "knowledge.zip_bomb"


def test_extractor_routes_json_end_to_end() -> None:
    extractor = DocumentContentExtractor()
    payload = json.dumps([{"a": 1, "b": 2}, {"a": 3, "b": 4}]).encode()

    doc = extractor.extract(data=payload, filename="data.json", content_type="application/json")

    assert doc.source_ext == ".json"
    assert doc.content_type == "application/json"
    assert doc.metadata["file_type"] == "structured_json"
    assert doc.metadata["parser"] == "json"
    assert len(doc.chunks) == 1
    assert doc.chunks[0].metadata["table_name"] == "data__table_0"


def test_extractor_counts_pdf_pages_not_blocks() -> None:
    """Block granularity (plan step 2) would otherwise report three blocks on
    two pages as a five-page document."""
    extractor = DocumentContentExtractor()
    data = _blocks_pdf(
        [
            [
                "Page one, block one, comfortably above the threshold.",
                "Page one, block two, also above the threshold.",
            ],
            ["Page two, block one, comfortably above the threshold."],
        ]
    )

    doc = extractor.extract(data=data, filename="report.pdf", content_type="application/pdf")

    assert doc.metadata["file_type"] == "pdf_text"
    assert doc.metadata["page_count"] == 2
    assert doc.metadata["block_count"] == 3
    assert len(doc.chunks) == 3


def test_extractor_runs_the_table_pass_before_the_text_pass() -> None:
    """The triple PDF path (plan step 3 / `P-07`): a page carrying a table and
    a paragraph yields BOTH kinds of chunk, and the table's rows are not
    indexed a second time as text."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "Quarterly headcount by department.", fontsize=11)
    page.insert_text((72, 118), "Figures are counted at quarter end.", fontsize=11)
    y = 250.0
    for row in _SALARY_ROWS:
        for x, cell in zip(_TABLE_COLUMN_X, row, strict=False):
            page.insert_text((x, y), cell, fontsize=11)
        y += 22
    data = doc.tobytes()
    doc.close()

    extracted = DocumentContentExtractor().extract(
        data=data, filename="report.pdf", content_type="application/pdf"
    )

    kinds = [c.kind for c in extracted.chunks]
    assert ParsedChunkKind.TABLE in kinds
    assert ParsedChunkKind.TEXT in kinds
    assert extracted.metadata["file_type"] == "pdf_mixed"
    assert extracted.metadata["parser"] == "pdf_multipass"
    assert extracted.metadata["table_count"] == 1
    # Tables lead the list, and the table's own rows appear only there.
    assert extracted.chunks[0].kind is ParsedChunkKind.TABLE
    text_blob = " ".join(c.text for c in extracted.chunks if c.kind is ParsedChunkKind.TEXT)
    assert "Finance" not in text_blob
    assert "headcount" in text_blob


def test_extractor_reports_a_pdf_with_no_tables_as_text_only() -> None:
    extracted = DocumentContentExtractor().extract(
        data=_blocks_pdf([["A page of ordinary prose with no table on it at all."]]),
        filename="prose.pdf",
        content_type="application/pdf",
    )

    assert extracted.metadata["file_type"] == "pdf_text"
    assert extracted.metadata["table_count"] == 0
    assert extracted.metadata["block_count"] == 1


def test_extractor_routes_docx_end_to_end() -> None:
    """Plan step 4 (`P-08`): a DOCX carrying prose and a table reports both
    counts and classifies on their ratio."""

    def build(document: WordDocumentObject) -> None:
        document.add_heading("Attendance Policy", level=1)
        document.add_paragraph("Employees must record arrival and departure daily.")
        _add_table(document, [["Grade", "Allowance"], ["A1", "1500"]])

    extracted = DocumentContentExtractor().extract(
        data=_docx_bytes(build),
        filename="policy.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert extracted.source_ext == ".docx"
    assert extracted.metadata["parser"] == "docx"
    assert extracted.metadata["file_type"] == "docx_semi_structured_text"
    assert extracted.metadata["block_count"] == 1
    assert extracted.metadata["table_count"] == 1
    # `_enrich` names the table, as it does for every other TABLE-kind chunk.
    table = next(c for c in extracted.chunks if c.kind is ParsedChunkKind.TABLE)
    assert table.metadata["table_name"] == "policy__table_0"
    assert all(c.metadata["source_ext"] == ".docx" for c in extracted.chunks)


def test_the_docx_upload_whitelist_matches_the_docx_route() -> None:
    """A parser no upload can reach is not a parser. `Limits.allowed_mime`
    dropped the DOCX type while `_ROUTES` had no `.docx` entry (an accepted
    upload that could never be parsed); plan step 4 (`P-08`) closes the pair
    from both ends, and this pins them together."""
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    assert mime in Limits().allowed_mime
    assert DocumentContentExtractor().supports(content_type=mime, filename="policy.docx") is True


def test_extractor_raises_validation_error_on_a_broken_docx() -> None:
    with pytest.raises(ValidationError):
        DocumentContentExtractor().extract(
            data=b"not a word document at all",
            filename="broken.docx",
            content_type="application/octet-stream",
        )


def test_extractor_routes_text_end_to_end() -> None:
    extractor = DocumentContentExtractor()
    doc = extractor.extract(data=b"hello world", filename="notes.txt", content_type="text/plain")

    assert doc.source_ext == ".txt"
    assert doc.metadata["parser"] == "text_plain"
    assert doc.chunks[0].metadata["source_ext"] == ".txt"


def test_content_extractor_protocol_conformance() -> None:
    # Structural (Protocol) conformance, verified statically by mypy and
    # exercised here at runtime.
    extractor: ContentExtractor = DocumentContentExtractor()
    assert extractor.supports(content_type="application/pdf", filename="a.pdf") is True


def test_extractor_reports_the_ocr_pass_on_a_docx(monkeypatch: pytest.MonkeyPatch) -> None:
    """The image pass joins the DOCX route's own chunks, and its two facts —
    how many images were read, and whether a cap cut it short — are reported
    in-band like every other route's counters."""
    _stub_ocr(monkeypatch)

    def build(document: WordDocumentObject) -> None:
        document.add_paragraph("Employees must clock in before 08:00.")
        document.add_picture(io.BytesIO(_png(400, 400)))

    parsed = DocumentContentExtractor().extract(
        data=_docx_bytes(build),
        filename="policy.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert parsed.metadata["image_count"] == 1
    assert parsed.metadata["ocr_truncated"] is False
    assert [chunk.kind for chunk in parsed.chunks] == [ParsedChunkKind.TEXT, ParsedChunkKind.OCR]
    assert parsed.chunks[-1].metadata["source_ext"] == ".docx"


def test_extractor_passes_its_limits_down_to_the_ocr_pass(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """`Settings` numbers that never reach a parser are decoration; this is the
    seam the worker's composition root uses (and step 14 will reuse)."""
    _stub_ocr(monkeypatch)
    data = _image_docx([_png(400, 400, tint=10), _png(400, 400, tint=20)])

    extractor = DocumentContentExtractor(limits=Limits(ocr_max_images_per_document=1))
    parsed = extractor.extract(data=data, filename="many.docx", content_type="x")

    assert parsed.metadata["image_count"] == 1
    assert parsed.metadata["ocr_truncated"] is True


@pytest.mark.parametrize(
    ("filename", "builder", "expected"),
    [
        ("shots.xlsx", _image_xlsx, "excel_images"),
        ("shots.docx", _image_docx, "docx_images"),
    ],
)
def test_extractor_never_calls_an_image_only_office_file_empty(
    monkeypatch: pytest.MonkeyPatch,
    filename: str,
    builder: Callable[[list[bytes]], bytes],
    expected: str,
) -> None:
    """A workbook whose only content is a pasted screenshot HAS chunks; calling
    it `*_empty` would be a lie about the document that reaches storage."""
    _stub_ocr(monkeypatch)

    parsed = DocumentContentExtractor().extract(
        data=builder([_png(400, 400)]), filename=filename, content_type="x"
    )

    assert parsed.metadata["file_type"] == expected
    assert len(parsed.chunks) == 1


def test_extractor_classifies_a_scanned_pdf_as_pdf_images(monkeypatch: pytest.MonkeyPatch) -> None:
    """The one case images decide: a PDF with no extractable text and no table
    used to be reported `pdf_empty`, which is the one thing it never was."""
    _stub_ocr(monkeypatch)

    parsed = DocumentContentExtractor().extract(
        data=_image_pdf([[_png(400, 400)]]), filename="scan.pdf", content_type="application/pdf"
    )

    assert parsed.metadata["file_type"] == "pdf_images"
    assert parsed.metadata["page_count"] == 1


def test_classify_docx_calls_a_picture_only_document_images_not_empty() -> None:
    assert docx_parser.classify_docx(table_count=0, text_count=0, image_count=2) == "docx_images"
    assert docx_parser.classify_docx(table_count=0, text_count=0, image_count=0) == "docx_empty"


# --------------------------------------------------------------------------- #
# extractor._enrich -- `file_name` on every route (rag-indexing-plan.md      #
# §3.9/§4 step 11, ح-١٣/ح-١٦: `_CITATION_KEYS` allowlisted `file_name` since  #
# step 11, but no producer ever emitted it until `_enrich` was closed to     #
# stamp it -- this is that fix, proven for every real producer.)             #
# --------------------------------------------------------------------------- #
def test_extractor_stamps_file_name_on_every_chunk_across_every_text_route() -> None:
    """PDF text, PDF tables, DOCX text+table, Excel, and JSON all route
    through `extractor.py`'s `_enrich`, so ALL of them get `file_name`
    stamped on every chunk -- checked here in one pass rather than editing
    every existing route test individually."""
    extractor = DocumentContentExtractor()

    json_doc_ = extractor.extract(
        data=json.dumps([{"a": 1, "b": 2}, {"a": 3, "b": 4}]).encode(),
        filename="quarterly-report.json",
        content_type="application/json",
    )
    assert json_doc_.chunks
    assert all(c.metadata["file_name"] == "quarterly-report.json" for c in json_doc_.chunks)

    text_doc = extractor.extract(
        data=b"hello world", filename="notes.txt", content_type="text/plain"
    )
    assert text_doc.chunks
    assert all(c.metadata["file_name"] == "notes.txt" for c in text_doc.chunks)

    pdf_text_doc = extractor.extract(
        data=_blocks_pdf([["A page of ordinary prose with no table on it at all."]]),
        filename="prose.pdf",
        content_type="application/pdf",
    )
    assert pdf_text_doc.chunks
    assert all(c.metadata["file_name"] == "prose.pdf" for c in pdf_text_doc.chunks)

    pdf_tables_doc = extractor.extract(
        data=_table_pdf_bytes([_SALARY_ROWS]),
        filename="salaries.pdf",
        content_type="application/pdf",
    )
    assert pdf_tables_doc.chunks
    assert all(c.metadata["file_name"] == "salaries.pdf" for c in pdf_tables_doc.chunks)

    def build_docx(document: WordDocumentObject) -> None:
        document.add_heading("Attendance Policy", level=1)
        document.add_paragraph("Employees must record arrival and departure daily.")
        _add_table(document, [["Grade", "Allowance"], ["A1", "1500"]])

    docx_doc = extractor.extract(
        data=_docx_bytes(build_docx),
        filename="policy.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert docx_doc.chunks
    assert all(c.metadata["file_name"] == "policy.docx" for c in docx_doc.chunks)
    # A file's TABLE-kind chunk carries BOTH `file_name` and `table_name` --
    # the two keys never collide or overwrite each other.
    table = next(c for c in docx_doc.chunks if c.kind is ParsedChunkKind.TABLE)
    assert table.metadata["table_name"] == "policy__table_0"

    def build_workbook(wb: openpyxl.Workbook) -> None:
        ws = wb.active
        ws.title = "Data"
        ws.append(["ID", "Value"])
        for i in range(3):
            ws.append([i, i * 2])

    excel_doc = extractor.extract(
        data=_workbook_bytes(build_workbook),
        filename="ledger.xlsx",
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert excel_doc.chunks
    assert all(c.metadata["file_name"] == "ledger.xlsx" for c in excel_doc.chunks)


def test_extractor_stamps_file_name_on_ocr_chunks_too(monkeypatch: pytest.MonkeyPatch) -> None:
    """The OCR pass (plan step 5) joins the DOCX route's own chunk list, and
    it is NOT exempt from `_enrich` -- its chunk gets `file_name` exactly
    like the text chunk next to it."""
    _stub_ocr(monkeypatch)

    def build(document: WordDocumentObject) -> None:
        document.add_paragraph("Employees must clock in before 08:00.")
        document.add_picture(io.BytesIO(_png(400, 400)))

    parsed = DocumentContentExtractor().extract(
        data=_docx_bytes(build),
        filename="policy-with-photo.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    ocr_chunk = next(c for c in parsed.chunks if c.kind is ParsedChunkKind.OCR)
    assert ocr_chunk.metadata["file_name"] == "policy-with-photo.docx"


# --------------------------------------------------------------------------- #
# extractor -> application.indexing.IndexDocument -- `file_name` genuinely   #
# reaches the built Qdrant payload, not just the citation allowlist          #
# (rag-indexing-plan.md §3.9, ح-١٣: a filtered key absent from every point   #
# matches nothing, silently).                                                #
# --------------------------------------------------------------------------- #
class _ProbeEmbeddings:
    """The smallest possible `EmbeddingProvider`: fixed-dimension vectors,
    no real model call -- this test only cares whether `file_name` reaches
    the payload, not what the vector looks like."""

    provider = "probe"

    async def embed(self, texts: list[str], model: str, api_key: str) -> EmbeddingResult:
        return EmbeddingResult(
            vectors=[[0.1, 0.2] for _ in texts], model=model, dimensions=2, tokens=len(texts)
        )

    def dimensions(self, model: str) -> int:
        return 2


class _ProbeVectors:
    """The smallest possible `HybridVectorStore`: records every upserted
    point's payload for inspection, nothing else."""

    def __init__(self) -> None:
        self.points: dict[str, VectorPoint] = {}

    async def ensure_hybrid_collection(
        self, name: str, dim: int, *, distance: str = "cosine"
    ) -> None: ...

    async def upsert(self, collection: str, points: list[VectorPoint]) -> None:
        for point in points:
            self.points[point.id] = point


async def test_extractor_output_file_name_reaches_the_built_qdrant_payload() -> None:
    """The full wiring, end to end: `DocumentContentExtractor.extract` (the
    real adapter, not a hand-built `ParsedChunk`) feeds `IndexDocument.
    execute`, and the resulting Qdrant point payload carries the CORRECT
    `file_name` -- proof it is genuinely present, not merely allowlisted in
    `_CITATION_KEYS`."""
    parsed: ParsedDocument = DocumentContentExtractor().extract(
        data=b"Quarterly headcount and revenue figures for the finance team.",
        filename="finance-quarterly-notes.txt",
        content_type="text/plain",
    )
    embeddings = _ProbeEmbeddings()
    vectors = _ProbeVectors()
    ctx = ExecutionContext(
        workspace_id="ws1", user_id="u1", correlation_id="corr", roles=frozenset({"member"})
    )

    outcome = await IndexDocument(embeddings, vectors).execute(
        ctx, document_id="doc-1", space_id=None, parsed=parsed, model="m", api_key="k"
    )

    assert outcome.chunks
    point = vectors.points[outcome.chunks[0].chunk_id]
    payload: Json = point.payload
    assert payload["file_name"] == "finance-quarterly-notes.txt"
